import numpy as np 
import pandas as pd
import random
import pickle
import re
from docx import Document
from docx.enum.text import WD_BREAK
import math
import PyPDF2
from docx2pdf import convert
import os

def analyze_pdf(pdf_file):
    even_types = []

    # Открытие PDF файла для чтения
    with open(pdf_file, "rb") as file:
        reader = PyPDF2.PdfReader(file)

        for i in range(len(reader.pages)):
            page = reader.pages[i]
            text = page.extract_text()

            # Проверка наличия "типа" на странице и чётности номера страницы
            if "Type:" in text and (i+1) % 2 == 0:
                even_types.append(text[text.index("Type:"):].split("\n")[0])  # Добавление "типа" в список

    return even_types


def parse_questions(file_path, number_of_students, number_of_questions):

    # Открыть документ
    doc = Document(file_path)
    full_text = []

    # Извлечение текста из каждого параграфа
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Объединение всех параграфов в одну строку
    input_text =  '\n'.join(full_text)

    # Разбиваем текст на отдельные вопросы
    questions_split = input_text.strip().split('\n\n')
    questions_list = []

    new_doc = Document()

    for question in questions_split:
        # Извлекаем текст вопроса
        question_text = re.search(r'\d+\)(.*?)\n[A-D]', question).group(1).strip()

        # Извлекаем варианты ответов
        options = re.findall(r'[A-D]\)(.*?)\n', question)

        # Извлекаем номер правильного ответа
        correct_answer = int(re.search(r'Answer:(\d+)', question).group(1))

        # Добавляем вопрос и ответы в список
        questions_list.append({
            "MC question": question_text,
            "options": options,
            "correct answer": correct_answer
        })

    questions = questions_list
    diction = {}
    for k in range(number_of_students):
        random.seed(k)
        np.random.seed(k)
        print(f"Type:{k+1}")
    # Randomly select 3 questions
        selected_questions = random.sample(questions, number_of_questions)
        # Dictionary to store the correct answers
        correct_answers = {}
        correct_nums = {}
        # Shuffle the options for each question and enumerate the questions
        for i, question in enumerate(selected_questions, start=1):
            correct_answer = question['options'][question['correct answer']]
            # print(correct_answer)
            # random.shuffle(question['options'])
            truth = [i == correct_answer for i in question["options"] ]
            print(f"Question {i}: {question['MC question']}")
            correct_nums[i-1] = np.where(truth)[0][0]
            for j, option in enumerate(question['options'], start=1):
                print(f"{chr(64+j)}: {option}")
            print("\n")
            correct_answers[f"Question {i}"] = correct_answer
        # print("Correct Answers:", correct_answers)
        diction[k+1] = correct_nums
        # print(correct_nums)
        with open(file_path.replace('.docx', ''), "wb") as file:
            pickle.dump(diction, file)
        # print(diction)

        # Запись вопросов в новый документ
        new_doc.add_paragraph(f"Type:{k+1}")
        for i, question in enumerate(selected_questions, start=1):
            new_doc.add_paragraph(f"Question {i}: {question['MC question']}")
            for j, option in enumerate(question['options'], start=1):
                new_doc.add_paragraph(f"{chr(64+j)}: {option}")
            new_doc.add_paragraph("")  # Пустая строка между вопросами
        
        # Предполагается, что new_doc - это экземпляр Document
        number_of_lines = len(new_doc.paragraphs)
        pages = math.ceil(number_of_lines / 26)  # Вычисляем количество страниц

        # Проверяем чётность страниц
        if pages % 2 != 0:
            # Если нечётное, добавляем один разрыв страницы
            new_doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        else:
            # Если чётное, добавляем два разрыва страницы
            new_doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            
    
    file_path_compiled = file_path.replace(".docx", "")
    
    file_path_compiled += r'_COMPILED.docx'
    # Сохранение нового документа
    new_doc.save(file_path_compiled)
    for i in range(number_of_students):
        convert(file_path_compiled)
        types_to_break = analyze_pdf(file_path_compiled.replace(".docx", ".pdf"))
        print(types_to_break)
        for i, paragraph in enumerate(new_doc.paragraphs):
            try:
                type_ = types_to_break[0]
                if type_[:-2] in paragraph.text:
                    # Добавляем разрыв страницы перед параграфом с "типом"
                    previous_paragraph = new_doc.paragraphs[i-1] if i > 0 else None
                    if previous_paragraph:
                        previous_paragraph.add_run().add_break(WD_BREAK.PAGE)
            except:
                continue
        # Сохраняем изменения в документе
        new_doc.save(file_path_compiled)

    return diction

# diction = parse_questions(r"C:\Users\user\Desktop\test\questions.docx", number_of_students=10, number_of_questions=5)

def parse_students(file_path):
    # Чтение данных из Excel файла
    df = pd.read_excel(file_path)

    # Убедитесь, что имена колонок соответствуют вашим данным
    df = df[['Students', 'St id']]

    # Сортировка по колонке с именами студентов
    # df.sort_values(by='Students', inplace=True)

    # Создание словаря
    students_dict = {i+1: {row['Students']: row['St id']} for i, row in df.iterrows()}
    with open(file_path.replace('.xlsx', ''), "wb") as file:
            pickle.dump(students_dict, file)
    return students_dict

# Пример использования
# file_path = r"C:\Users\user\Desktop\students.xlsx"
# students_dict = parse_students(file_path)
# print(students_dict)

    
# file_path = r"C:\Users\user\Desktop\questions.docx"
# text = parse_questions(file_path)
# print(text)
